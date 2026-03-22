from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocxDocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.table import _Cell as DocxCell
from docx.text.paragraph import Paragraph as DocxParagraph

from papermentor_os.schemas.paper import PaperPackage, PaperReference, Paragraph, Section
from papermentor_os.schemas.types import Discipline, PaperStage
from papermentor_os.shared.text import normalize_whitespace


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WORD_NS}}}"


HEADING_STYLE_PATTERN = re.compile(r"heading\s*(\d+)", re.IGNORECASE)
NUMBERED_HEADING_PATTERN = re.compile(r"^\d+(?:\.\d+)*\s+")
SECTION_HEADING_PATTERN = re.compile(r"^(?:第[一二三四五六七八九十百]+章|[一二三四五六七八九十]+、|\d+(?:\.\d+)*\s+)")
HEADING_NUMBER_PATTERN = re.compile(r"^(?P<number>\d+(?:\.\d+)*)\s+(?P<label>.+)$")
ABSTRACT_TITLES = {"摘要", "abstract"}
KEYWORD_TITLES = {"关键词", "关键字", "keywords", "keyword"}
KEYWORD_LINE_PREFIXES = ("关键词", "关键字", "keywords", "keyword", "key words")
REFERENCE_TITLES = {"参考文献", "references", "bibliography"}
CONTENTS_TITLES = {"目录", "contents", "tableofcontents"}
APPENDIX_CONTENTS_TITLES = {
    "附录目录",
    "附录图目录",
    "附录表目录",
    "appendixcontents",
    "listofappendices",
    "listofappendixfigures",
    "listofappendixtables",
}
METADATA_LINE_PREFIXES = {
    "分类号",
    "中图分类号",
    "学校代码",
    "学号",
    "保密级别",
    "作者姓名",
    "作者",
    "指导教师",
    "学院",
    "院系",
    "专业",
    "班级",
    "作者单位",
    "导师组",
    "导师组成员",
    "学位类型",
    "学位类别",
    "学位授予单位",
    "udc",
    "studentid",
}
SKIPPABLE_FRONT_MATTER_TITLES = {
    "独创性声明",
    "原创性声明",
    "诚信声明",
    "学术诚信承诺书",
    "独创性承诺书",
    "版权声明",
}
BACK_MATTER_TITLES = {
    "致谢",
    "后记",
    "acknowledgements",
    "acknowledgments",
    "acknowledgement",
    "acknowledgment",
    "攻读学位期间取得的研究成果",
    "攻读学位期间取得的成果",
    "攻读学位期间发表的成果",
    "科研成果",
    "研究成果",
}
POST_REFERENCE_BACK_MATTER_TITLES = {
    "作者简介",
    "个人简历",
    "作者简历",
    "abouttheauthor",
    "authorbiography",
    "authorbio",
    "biographicalnote",
}
APPENDIX_HEADING_PATTERN = re.compile(
    r"^(?:附录|appendix)(?:[\s:：-]*[A-Z0-9一二三四五六七八九十]+)?(?:[\s:：-]+.+)?$",
    re.IGNORECASE,
)
TITLE_STYLE_HINTS = {"title", "标题", "论文题目"}
COVER_PAGE_HINTS = {"大学", "学院", "专业", "指导教师", "学号", "姓名", "班级"}
TITLE_KEYWORDS = {"评审", "论文", "研究", "设计", "实现", "框架", "分析", "系统", "方法", "适配"}
TOC_ENTRY_PATTERN = re.compile(
    r"^(?:第[一二三四五六七八九十百]+章|[一二三四五六七八九十]+、|附录[A-Z]?|appendix\s+[A-Z]?|\d+(?:\.\d+)*\.?(?:\s*|(?=[A-Za-z\u4e00-\u9fff]))).*(?:\.{2,}|\u2026{2,}|\s+|\t+)\d+$",
    re.IGNORECASE,
)
TOC_FIELD_CODE_ARTIFACT_PATTERN = re.compile(
    r"(?:\bPAGEREF\b|\bHYPERLINK\b|\bTOC\b|_Toc\d+|\\[hzup])",
    re.IGNORECASE,
)
CAPTION_PATTERN = re.compile(
    r"^(?:图|表|式|公式|figure|table|equation)\s*(?:\(?[A-Za-z]?\d+(?:[-.]\d+)*\)?)(?:[\s:：-]+.+)?$",
    re.IGNORECASE,
)
ANNOTATION_PATTERN = re.compile(
    r"^(?:注|注释|说明|备注|note|notes|remark|remarks)\s*(?:(?:[:：]\s*.+)|(?:\(?[A-Za-z]?\d+(?:[-.]\d+)*\)?(?:[\s:：-]+.+)?))$",
    re.IGNORECASE,
)
PAGE_NUMBER_NOISE_PATTERN = re.compile(
    r"^(?:[-–—]\s*)?(?:\d{1,3}|[ivxlcdm]{1,8})(?:\s*[-–—])?$",
    re.IGNORECASE,
)
FOOTNOTE_MARKER_PATTERN = re.compile(
    r"^(?:\[\d+\]|\(\d+\)|\d+\)|[①②③④⑤⑥⑦⑧⑨⑩]+|\*+)$"
)
FOOTNOTE_BODY_PATTERN = re.compile(
    r"^(?:\[\d+\]|\(\d+\)|\d+\)|[①②③④⑤⑥⑦⑧⑨⑩]+)\s*(?:(?:注|注释|说明|备注|note|remark|remarks)[:：]?\s*.+|.+)$",
    re.IGNORECASE,
)
RUNNING_HEADER_FOOTER_HINTS = {
    "本科毕业论文",
    "学士学位论文",
    "毕业设计（论文）",
}
RUNNING_ENGLISH_HEADER_FOOTER_HINTS = {
    "undergraduate thesis",
    "bachelor thesis",
    "bachelor's thesis",
}
RUNNING_ENGLISH_METADATA_PATTERN = re.compile(
    r"^(?:student\s*id|advisor|supervisor|major|department|school)\b[\s:：-]+.+$",
    re.IGNORECASE,
)


class DocxPaperParser:
    """Parse V1 docx thesis drafts into a stable PaperPackage."""

    def parse_file(
        self,
        file_path: str | Path,
        *,
        stage: PaperStage = PaperStage.DRAFT,
        discipline: Discipline = Discipline.COMPUTER_SCIENCE,
    ) -> PaperPackage:
        path = Path(file_path)
        if path.suffix.lower() != ".docx":
            raise ValueError("V1 only supports .docx input.")

        document = Document(str(path))
        blocks = self._extract_block_items(document)
        if not blocks:
            raise ValueError("The docx file does not contain readable paragraphs.")

        title_index = self._select_title_index(blocks)
        title = blocks[title_index]["text"]
        body_items = blocks[title_index + 1:]
        note_texts = self._extract_note_texts(path)

        abstract_lines: list[str] = []
        sections: list[Section] = []
        references: list[PaperReference] = []

        current_section: Section | None = None
        mode = "body"
        body_started = False
        section_counter = 0
        paragraph_counter = 0
        reference_counter = 0

        for item in body_items:
            text = item["text"]
            if self._is_structural_noise_line(text):
                continue
            if mode == "body" and body_started and self._is_docx_note_body(text, note_texts):
                continue
            if mode == "body" and body_started and self._is_running_footnote_body_noise(text):
                continue
            if body_started and self._is_running_header_footer_noise(text, title):
                continue
            if current_section is not None and self._is_repeated_running_section_heading(text, current_section.heading):
                continue
            heading_level = self._heading_level(text, item["style"], item.get("kind", "paragraph"))
            if self._is_repeated_seen_section_heading(text, heading_level, current_section, sections):
                continue
            normalized_heading = self._canonical_heading_label(text)

            if mode in {"abstract", "keywords"}:
                if normalized_heading in ABSTRACT_TITLES:
                    mode = "abstract"
                    current_section = None
                    continue
                if self._is_keyword_line(text):
                    abstract_lines.append(text)
                    mode = "keywords"
                    continue
                if self._is_metadata_line(text):
                    mode = "front_matter"
                    continue
                if normalized_heading in CONTENTS_TITLES:
                    mode = "contents"
                    current_section = None
                    continue
                if normalized_heading in SKIPPABLE_FRONT_MATTER_TITLES:
                    mode = "front_matter"
                    current_section = None
                    continue
                if not self._looks_like_section_start(text, item["style"]):
                    abstract_lines.append(text)
                    continue

            if mode == "front_matter":
                if normalized_heading in ABSTRACT_TITLES:
                    mode = "abstract"
                    current_section = None
                    continue
                if self._is_keyword_line(text):
                    abstract_lines.append(text)
                    mode = "keywords"
                    continue
                if normalized_heading in CONTENTS_TITLES:
                    mode = "contents"
                    current_section = None
                    continue
                if self._is_metadata_line(text) or normalized_heading in SKIPPABLE_FRONT_MATTER_TITLES:
                    continue
                if not self._looks_like_section_start(text, item["style"]):
                    continue

            if mode == "appendix" and normalized_heading not in REFERENCE_TITLES:
                continue

            if mode == "back_matter":
                if normalized_heading in REFERENCE_TITLES:
                    mode = "references"
                    current_section = None
                    continue
                if heading_level is not None and normalized_heading in BACK_MATTER_TITLES:
                    continue
                continue

            if mode == "post_references_back_matter":
                continue

            if mode == "contents" and (
                self._is_table_of_contents_entry(text) or self._is_table_of_contents_artifact(text)
            ):
                continue
            if mode == "contents" and not self._looks_like_section_start(text, item["style"]):
                continue

            if heading_level is not None:
                if normalized_heading in ABSTRACT_TITLES:
                    mode = "abstract"
                    current_section = None
                    continue
                if normalized_heading in CONTENTS_TITLES:
                    mode = "contents"
                    current_section = None
                    continue
                if not body_started and normalized_heading in KEYWORD_TITLES:
                    mode = "keywords"
                    current_section = None
                    continue
                if not body_started and normalized_heading in SKIPPABLE_FRONT_MATTER_TITLES:
                    mode = "front_matter"
                    current_section = None
                    continue
                if normalized_heading in REFERENCE_TITLES:
                    mode = "references"
                    current_section = None
                    continue
                if mode == "references" and normalized_heading in POST_REFERENCE_BACK_MATTER_TITLES:
                    mode = "post_references_back_matter"
                    current_section = None
                    continue
                if body_started and normalized_heading in BACK_MATTER_TITLES:
                    mode = "back_matter"
                    current_section = None
                    continue
                if body_started and normalized_heading in APPENDIX_CONTENTS_TITLES:
                    mode = "appendix"
                    current_section = None
                    continue
                if body_started and self._is_appendix_heading(text):
                    mode = "appendix"
                    current_section = None
                    continue

                mode = "body"
                body_started = True
                section_counter += 1
                current_section = Section(
                    section_id=f"sec-{section_counter:03d}",
                    heading=text,
                    level=heading_level,
                    paragraphs=[],
                )
                sections.append(current_section)
                continue

            if mode == "abstract":
                abstract_lines.append(text)
                continue

            if mode == "references":
                if heading_level is not None and normalized_heading in POST_REFERENCE_BACK_MATTER_TITLES:
                    mode = "post_references_back_matter"
                    current_section = None
                    continue
                reference_counter += 1
                references.append(
                    PaperReference(
                        reference_id=f"ref-{reference_counter:03d}",
                        anchor_id=f"ref-{reference_counter:03d}",
                        raw_text=text,
                    )
                )
                continue

            if mode in {"contents", "front_matter", "appendix", "back_matter", "post_references_back_matter"}:
                continue

            if current_section is None:
                body_started = True
                section_counter += 1
                current_section = Section(
                    section_id=f"sec-{section_counter:03d}",
                    heading="未命名章节",
                    level=1,
                    paragraphs=[],
                )
                sections.append(current_section)

            paragraph_counter += 1
            current_section.paragraphs.append(
                Paragraph(
                    paragraph_id=f"p-{paragraph_counter:04d}",
                    anchor_id=f"{current_section.section_id}-p-{len(current_section.paragraphs) + 1:03d}",
                    text=text,
                )
            )

        return PaperPackage(
            paper_id=path.stem,
            title=title,
            discipline=discipline,
            stage=stage,
            abstract="\n".join(abstract_lines),
            sections=sections,
            references=references,
            source_path=str(path),
        )

    def _heading_level(self, text: str, style_name: str, kind: str = "paragraph") -> int | None:
        normalized_heading = self._canonical_heading_label(text)
        if style_name:
            matched = HEADING_STYLE_PATTERN.search(style_name)
            if matched:
                return max(1, min(6, int(matched.group(1))))
            if style_name in {"标题 1", "标题 2", "标题 3"}:
                return int(style_name.split()[-1])

        if normalized_heading in ABSTRACT_TITLES | REFERENCE_TITLES:
            return 1

        if kind == "table":
            return None

        if NUMBERED_HEADING_PATTERN.match(text):
            return min(text.count(".") + 1, 6)

        if self._is_figure_or_table_caption(text):
            return None
        if self._is_annotation_block(text):
            return None

        if len(text) <= 30 and not text.endswith(("。", ".", "；", ";")):
            return 1

        return None

    def _canonical_heading_label(self, text: str) -> str:
        return re.sub(r"\s+", "", text).lower()

    def _is_table_of_contents_entry(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        return bool(TOC_ENTRY_PATTERN.match(normalized))

    def _is_table_of_contents_artifact(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        return bool(TOC_FIELD_CODE_ARTIFACT_PATTERN.search(normalized))

    def _is_keyword_line(self, text: str) -> bool:
        normalized = self._canonical_heading_label(text)
        if normalized in KEYWORD_TITLES:
            return True
        return self._matches_prefixed_value_line(text, KEYWORD_LINE_PREFIXES)

    def _is_metadata_line(self, text: str) -> bool:
        return self._matches_prefixed_value_line(text, tuple(METADATA_LINE_PREFIXES))

    def _is_appendix_heading(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        return bool(APPENDIX_HEADING_PATTERN.match(normalized))

    def _is_figure_or_table_caption(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        return bool(CAPTION_PATTERN.match(normalized))

    def _is_annotation_block(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        return bool(ANNOTATION_PATTERN.match(normalized))

    def _is_structural_noise_line(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        return bool(
            PAGE_NUMBER_NOISE_PATTERN.match(normalized)
            or FOOTNOTE_MARKER_PATTERN.match(normalized)
        )

    def _is_running_footnote_body_noise(self, text: str) -> bool:
        normalized = normalize_whitespace(text)
        if not FOOTNOTE_BODY_PATTERN.match(normalized):
            return False
        return len(normalized) <= 80

    def _is_docx_note_body(self, text: str, note_texts: set[str]) -> bool:
        normalized = normalize_whitespace(text)
        return len(normalized) <= 120 and normalized in note_texts

    def _is_running_header_footer_noise(self, text: str, title: str) -> bool:
        normalized = normalize_whitespace(text)
        canonical = self._canonical_heading_label(normalized)
        if canonical == self._canonical_heading_label(title):
            return True
        if canonical in ABSTRACT_TITLES:
            return True
        if any(hint in normalized for hint in RUNNING_HEADER_FOOTER_HINTS) and len(normalized) <= 30:
            return True
        lowered = normalized.lower()
        if any(hint in lowered for hint in RUNNING_ENGLISH_HEADER_FOOTER_HINTS) and len(normalized) <= 40:
            return True
        if self._is_metadata_line(normalized):
            return True
        if RUNNING_ENGLISH_METADATA_PATTERN.match(normalized) and len(normalized) <= 40:
            return True
        hint_hits = sum(1 for hint in COVER_PAGE_HINTS if hint in normalized)
        if hint_hits >= 2 and len(normalized) <= 30 and not normalized.endswith(("。", ".", "；", ";", "：", ":")):
            return True
        return False

    def _is_repeated_running_section_heading(self, text: str, current_heading: str) -> bool:
        normalized = normalize_whitespace(text)
        if not normalized:
            return False
        return self._canonical_heading_label(normalized) == self._canonical_heading_label(current_heading)

    def _is_repeated_seen_section_heading(
        self,
        text: str,
        heading_level: int | None,
        current_section: Section | None,
        sections: list[Section],
    ) -> bool:
        if heading_level is None or current_section is None:
            return False

        canonical = self._canonical_heading_label(text)
        if not canonical:
            return False

        return any(
            self._is_eligible_running_heading_target(section, current_section, heading_level)
            and (
                self._canonical_heading_label(section.heading) == canonical
                or self._is_abbreviated_heading_variant(text, section.heading)
                or self._is_unnumbered_abbreviated_heading_variant(text, section.heading)
            )
            for section in sections
        )

    def _is_eligible_running_heading_target(
        self,
        section: Section,
        current_section: Section,
        heading_level: int,
    ) -> bool:
        return section.level <= current_section.level and section.level <= heading_level

    def _is_abbreviated_heading_variant(self, text: str, section_heading: str) -> bool:
        candidate_number, candidate_label = self._split_numbered_heading(text)
        section_number, section_label = self._split_numbered_heading(section_heading)
        if candidate_number is None or section_number is None:
            return False
        if candidate_number != section_number:
            return False

        candidate_canonical = self._canonical_heading_label(candidate_label)
        section_canonical = self._canonical_heading_label(section_label)
        if not candidate_canonical or not section_canonical:
            return False
        if candidate_canonical == section_canonical:
            return False
        if len(candidate_canonical) < 4:
            return False
        return (
            len(candidate_canonical) < len(section_canonical)
            and section_canonical.startswith(candidate_canonical)
        )

    def _is_unnumbered_abbreviated_heading_variant(self, text: str, section_heading: str) -> bool:
        candidate_number, candidate_label = self._split_numbered_heading(text)
        section_number, section_label = self._split_numbered_heading(section_heading)
        if candidate_number is not None or section_number is None:
            return False

        candidate_canonical = self._canonical_heading_label(candidate_label)
        section_canonical = self._canonical_heading_label(section_label)
        if not candidate_canonical or not section_canonical:
            return False
        if candidate_canonical == section_canonical:
            return False
        if len(candidate_canonical) < 4:
            return False
        return (
            len(candidate_canonical) < len(section_canonical)
            and section_canonical.startswith(candidate_canonical)
        )

    def _split_numbered_heading(self, text: str) -> tuple[str | None, str]:
        normalized = normalize_whitespace(text)
        matched = HEADING_NUMBER_PATTERN.match(normalized)
        if matched is None:
            return None, normalized
        return matched.group("number"), matched.group("label")

    def _looks_like_section_start(self, text: str, style_name: str) -> bool:
        if style_name:
            matched = HEADING_STYLE_PATTERN.search(style_name)
            if matched:
                return True
            if style_name in {"标题 1", "标题 2", "标题 3"}:
                return True
        return bool(SECTION_HEADING_PATTERN.match(text))

    def _extract_block_items(self, document: DocxDocumentObject) -> list[dict[str, str]]:
        items: list[dict[str, str]] = []
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = DocxParagraph(child, document)
                text = normalize_whitespace(paragraph.text)
                if not text:
                    continue
                items.append(
                    {
                        "text": text,
                        "style": normalize_whitespace(getattr(paragraph.style, "name", "")),
                        "kind": "paragraph",
                    }
                )
                continue

            if isinstance(child, CT_Tbl):
                table = DocxTable(child, document)
                items.extend(self._extract_table_rows(table))

        return items

    def _extract_table_rows(self, table: DocxTable) -> list[dict[str, str]]:
        rows: list[dict[str, str]] = []
        for row in table.rows:
            cells = [self._extract_cell_text(cell) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            cells = self._dedupe_adjacent_values(cells)
            if not cells:
                continue
            rows.append({"text": " ".join(cells), "style": "", "kind": "table"})
        return rows

    def _extract_cell_text(self, cell: DocxCell) -> str:
        parts: list[str] = []
        for child in cell._tc.iterchildren():
            if isinstance(child, CT_P):
                paragraph = DocxParagraph(child, cell)
                text = normalize_whitespace(paragraph.text)
                if text:
                    parts.append(text)
                continue

            if isinstance(child, CT_Tbl):
                nested_table = DocxTable(child, cell)
                nested_rows = self._extract_table_rows(nested_table)
                parts.extend(row["text"] for row in nested_rows if row["text"])

        return " ".join(self._dedupe_adjacent_values(parts))

    def _dedupe_adjacent_values(self, values: list[str]) -> list[str]:
        deduped: list[str] = []
        for value in values:
            if deduped and deduped[-1] == value:
                continue
            deduped.append(value)
        return deduped

    def _matches_prefixed_value_line(self, text: str, prefixes: tuple[str, ...]) -> bool:
        normalized = normalize_whitespace(text).strip()
        lowered = normalized.lower()
        for prefix in sorted(prefixes, key=len, reverse=True):
            lowered_prefix = prefix.lower()
            if lowered == lowered_prefix:
                continue
            if not lowered.startswith(lowered_prefix):
                continue

            remainder = normalized[len(prefix) :].lstrip()
            if not remainder:
                continue

            if remainder[0] in {":", "：", "-", "–", "—"}:
                return len(remainder[1:].strip()) > 0
            return True

        return False

    def _select_title_index(self, paragraphs: list[dict[str, str]]) -> int:
        abstract_index = next(
            (
                index
                for index, item in enumerate(paragraphs)
                if self._canonical_heading_label(item["text"]) in ABSTRACT_TITLES
            ),
            len(paragraphs),
        )
        candidate_range = paragraphs[:abstract_index] or paragraphs[:1]

        scored_candidates: list[tuple[int, int]] = []
        for index, item in enumerate(candidate_range):
            text = item["text"]
            style_name = item["style"]
            score = 0
            canonical_style = style_name.lower()
            if any(hint in canonical_style for hint in TITLE_STYLE_HINTS):
                score += 6
            if any(keyword in text for keyword in TITLE_KEYWORDS):
                score += 2
            if 8 <= len(text) <= 40:
                score += 2
            if len(text) > 45:
                score -= 2
            if any(hint in text for hint in COVER_PAGE_HINTS):
                score -= 4
            if self._is_metadata_line(text) or self._is_keyword_line(text):
                score -= 6
            if self._is_table_of_contents_entry(text) or self._is_table_of_contents_artifact(text):
                score -= 8
            if self._canonical_heading_label(text) in ABSTRACT_TITLES | REFERENCE_TITLES:
                score -= 10
            scored_candidates.append((score, index))

        best_score, best_index = max(scored_candidates, key=lambda item: (item[0], -item[1]))
        if best_score <= 0:
            return 0
        return best_index

    def _extract_note_texts(self, path: Path) -> set[str]:
        try:
            with zipfile.ZipFile(path) as archive:
                note_texts = set()
                note_texts.update(self._extract_note_part_texts(archive, "word/footnotes.xml", "footnote"))
                note_texts.update(self._extract_note_part_texts(archive, "word/endnotes.xml", "endnote"))
                return note_texts
        except (OSError, zipfile.BadZipFile, ET.ParseError):
            return set()

    def _extract_note_part_texts(
        self,
        archive: zipfile.ZipFile,
        part_name: str,
        note_tag: str,
    ) -> set[str]:
        if part_name not in archive.namelist():
            return set()

        root = ET.fromstring(archive.read(part_name))
        note_texts: set[str] = set()
        for note in root.findall(f"{W}{note_tag}"):
            note_id = note.get(f"{W}id")
            if note_id in {"-1", "0"}:
                continue
            full_text = normalize_whitespace(" ".join(node.text or "" for node in note.iter(f"{W}t")))
            if full_text:
                note_texts.add(full_text)
            for paragraph in note.findall(f"{W}p"):
                paragraph_text = normalize_whitespace(" ".join(node.text or "" for node in paragraph.iter(f"{W}t")))
                if paragraph_text:
                    note_texts.add(paragraph_text)
        return note_texts
