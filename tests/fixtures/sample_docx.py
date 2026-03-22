from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocxDocumentObject
from docx.table import _Cell
from tests.fixtures.review_cases import MINIMAL_REVIEW_CASE, ParagraphSpec, ReviewCaseSpec, TableCellSpec, TableSpec


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
DOC_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
FOOTNOTE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
FOOTNOTE_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
ENDNOTE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
ENDNOTE_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
W = f"{{{WORD_NS}}}"
R = f"{{{REL_NS}}}"
CT = f"{{{PACKAGE_NS}}}"
PR = f"{{{DOC_REL_NS}}}"

ET.register_namespace("w", WORD_NS)
ET.register_namespace("r", REL_NS)


def build_docx_from_case(file_path: Path, case: ReviewCaseSpec) -> Path:
    document = Document()
    pending_footnotes: list[tuple[str, str | tuple[str, ...]]] = []
    pending_endnotes: list[tuple[str, str | tuple[str, ...]]] = []
    for front_matter_paragraph in case.front_matter:
        document.add_paragraph(front_matter_paragraph)
    for table_spec in case.front_matter_tables:
        _append_table(document, table_spec, pending_footnotes, pending_endnotes)
    if case.title_style is not None:
        document.add_paragraph(case.title, style=case.title_style)
    else:
        document.add_paragraph(case.title)
    document.add_paragraph(case.abstract_heading, style=case.heading_style)
    document.add_paragraph(case.abstract)
    for front_matter_paragraph in case.post_abstract_front_matter:
        document.add_paragraph(front_matter_paragraph)
    for table_spec in case.post_abstract_front_matter_tables:
        _append_table(document, table_spec, pending_footnotes, pending_endnotes)
    for section in case.sections:
        document.add_paragraph(section.heading, style=case.heading_style)
        for paragraph in section.paragraphs:
            paragraph_text = paragraph.text if isinstance(paragraph, ParagraphSpec) else paragraph
            docx_paragraph = document.add_paragraph(paragraph_text)
            if isinstance(paragraph, ParagraphSpec) and paragraph.footnote:
                placeholder = f"PM_FOOTNOTE_PLACEHOLDER_{len(pending_footnotes) + 1}"
                docx_paragraph.add_run(placeholder)
                pending_footnotes.append((placeholder, paragraph.footnote))
            if isinstance(paragraph, ParagraphSpec) and paragraph.endnote:
                placeholder = f"PM_ENDNOTE_PLACEHOLDER_{len(pending_endnotes) + 1}"
                docx_paragraph.add_run(placeholder)
                pending_endnotes.append((placeholder, paragraph.endnote))
        for table_spec in section.tables:
            _append_table(document, table_spec, pending_footnotes, pending_endnotes)
    if case.references:
        document.add_paragraph(case.reference_heading, style=case.heading_style)
        for reference in case.references:
            document.add_paragraph(reference)
    for section in case.post_reference_back_matter_sections:
        document.add_paragraph(section.heading, style=case.heading_style)
        for paragraph in section.paragraphs:
            paragraph_text = paragraph.text if isinstance(paragraph, ParagraphSpec) else paragraph
            document.add_paragraph(paragraph_text)
    document.save(file_path)
    if pending_footnotes:
        _inject_footnotes(file_path, pending_footnotes)
    if pending_endnotes:
        _inject_endnotes(file_path, pending_endnotes)
    return file_path


def build_minimal_review_docx(file_path: Path) -> Path:
    return build_docx_from_case(file_path, MINIMAL_REVIEW_CASE)


def _append_table(
    container: DocxDocumentObject | _Cell,
    table_spec: TableSpec | tuple[tuple[str, ...], ...],
    pending_footnotes: list[tuple[str, str | tuple[str, ...]]] | None = None,
    pending_endnotes: list[tuple[str, str | tuple[str, ...]]] | None = None,
) -> None:
    normalized_spec = _normalize_table_spec(table_spec)
    if not normalized_spec.rows:
        return

    column_count = max(len(row) for row in normalized_spec.rows)
    table = container.add_table(rows=0, cols=column_count)
    for row_cells in normalized_spec.rows:
        row = table.add_row().cells
        for index, cell_spec in enumerate(row_cells):
            _populate_cell(row[index], cell_spec, pending_footnotes, pending_endnotes)

    for start_row, start_col, end_row, end_col in normalized_spec.merges:
        table.cell(start_row, start_col).merge(table.cell(end_row, end_col))


def _normalize_table_spec(table_spec: TableSpec | tuple[tuple[str, ...], ...]) -> TableSpec:
    if isinstance(table_spec, TableSpec):
        return table_spec
    return TableSpec(rows=tuple(tuple(cell for cell in row) for row in table_spec))


def _populate_cell(
    cell: _Cell,
    cell_spec: TableCellSpec | str,
    pending_footnotes: list[tuple[str, str | tuple[str, ...]]] | None = None,
    pending_endnotes: list[tuple[str, str | tuple[str, ...]]] | None = None,
) -> None:
    if isinstance(cell_spec, TableCellSpec):
        if cell_spec.text or cell_spec.footnote or cell_spec.endnote:
            paragraph = cell.paragraphs[0]
            if cell_spec.text:
                paragraph.text = cell_spec.text
            if cell_spec.footnote and pending_footnotes is not None:
                placeholder = f"PM_FOOTNOTE_PLACEHOLDER_{len(pending_footnotes) + 1}"
                paragraph.add_run(placeholder)
                pending_footnotes.append((placeholder, cell_spec.footnote))
            if cell_spec.endnote and pending_endnotes is not None:
                placeholder = f"PM_ENDNOTE_PLACEHOLDER_{len(pending_endnotes) + 1}"
                paragraph.add_run(placeholder)
                pending_endnotes.append((placeholder, cell_spec.endnote))
        if cell_spec.nested_table_rows:
            _append_table(cell, cell_spec.nested_table_rows, pending_footnotes, pending_endnotes)
        return

    if cell_spec:
        cell.text = cell_spec


def _inject_footnotes(file_path: Path, pending_footnotes: list[tuple[str, str | tuple[str, ...]]]) -> None:
    with zipfile.ZipFile(file_path, "r") as source_zip:
        files = {name: source_zip.read(name) for name in source_zip.namelist()}

    files["[Content_Types].xml"] = _update_content_types(
        files["[Content_Types].xml"],
        "/word/footnotes.xml",
        FOOTNOTE_CONTENT_TYPE,
    )
    files["word/_rels/document.xml.rels"] = _update_document_relationships(
        files["word/_rels/document.xml.rels"],
        FOOTNOTE_REL_TYPE,
        "footnotes.xml",
        "rIdPmFootnotes",
    )
    files["word/document.xml"] = _update_document_with_notes(
        files["word/document.xml"],
        pending_footnotes,
        "footnoteReference",
    )
    files["word/footnotes.xml"] = _build_footnotes_part(pending_footnotes)

    with zipfile.ZipFile(file_path, "w") as target_zip:
        for name, content in files.items():
            target_zip.writestr(name, content)


def _inject_endnotes(file_path: Path, pending_endnotes: list[tuple[str, str | tuple[str, ...]]]) -> None:
    with zipfile.ZipFile(file_path, "r") as source_zip:
        files = {name: source_zip.read(name) for name in source_zip.namelist()}

    files["[Content_Types].xml"] = _update_content_types(
        files["[Content_Types].xml"],
        "/word/endnotes.xml",
        ENDNOTE_CONTENT_TYPE,
    )
    files["word/_rels/document.xml.rels"] = _update_document_relationships(
        files["word/_rels/document.xml.rels"],
        ENDNOTE_REL_TYPE,
        "endnotes.xml",
        "rIdPmEndnotes",
    )
    files["word/document.xml"] = _update_document_with_notes(
        files["word/document.xml"],
        pending_endnotes,
        "endnoteReference",
    )
    files["word/endnotes.xml"] = _build_endnotes_part(pending_endnotes)

    with zipfile.ZipFile(file_path, "w") as target_zip:
        for name, content in files.items():
            target_zip.writestr(name, content)


def _update_content_types(content: bytes, part_name: str, content_type: str) -> bytes:
    root = ET.fromstring(content)
    if not any(
        override.get("PartName") == part_name
        for override in root.findall(f"{CT}Override")
    ):
        ET.SubElement(
            root,
            f"{CT}Override",
            {
                "PartName": part_name,
                "ContentType": content_type,
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _update_document_relationships(content: bytes, rel_type: str, target: str, id_prefix: str) -> bytes:
    root = ET.fromstring(content)
    if not any(
        relationship.get("Type") == rel_type
        for relationship in root.findall(f"{PR}Relationship")
    ):
        existing_ids = {
            relationship.get("Id", "")
            for relationship in root.findall(f"{PR}Relationship")
        }
        next_index = 1
        relationship_id = f"{id_prefix}{next_index}"
        while relationship_id in existing_ids:
            next_index += 1
            relationship_id = f"{id_prefix}{next_index}"
        ET.SubElement(
            root,
            f"{PR}Relationship",
            {
                "Id": relationship_id,
                "Type": rel_type,
                "Target": target,
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _update_document_with_notes(
    content: bytes,
    pending_notes: list[tuple[str, str | tuple[str, ...]]],
    reference_tag: str,
) -> bytes:
    root = ET.fromstring(content)
    placeholder_to_id = {
        placeholder: str(index)
        for index, (placeholder, _text) in enumerate(pending_notes, start=1)
    }
    for run in root.iter(f"{W}r"):
        for text_element in list(run.findall(f"{W}t")):
            placeholder = text_element.text or ""
            note_id = placeholder_to_id.get(placeholder)
            if note_id is None:
                continue
            run.remove(text_element)
            run.append(ET.Element(f"{W}{reference_tag}", {f"{W}id": note_id}))
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _build_footnotes_part(pending_footnotes: list[tuple[str, str | tuple[str, ...]]]) -> bytes:
    root = ET.Element(f"{W}footnotes")
    separator = ET.SubElement(root, f"{W}footnote", {f"{W}type": "separator", f"{W}id": "-1"})
    separator_paragraph = ET.SubElement(separator, f"{W}p")
    separator_run = ET.SubElement(separator_paragraph, f"{W}r")
    ET.SubElement(separator_run, f"{W}separator")

    continuation = ET.SubElement(
        root,
        f"{W}footnote",
        {f"{W}type": "continuationSeparator", f"{W}id": "0"},
    )
    continuation_paragraph = ET.SubElement(continuation, f"{W}p")
    continuation_run = ET.SubElement(continuation_paragraph, f"{W}r")
    ET.SubElement(continuation_run, f"{W}continuationSeparator")

    for footnote_id, (_placeholder, text) in enumerate(pending_footnotes, start=1):
        footnote = ET.SubElement(root, f"{W}footnote", {f"{W}id": str(footnote_id)})
        for paragraph_text in _normalize_note_text(text):
            paragraph = ET.SubElement(footnote, f"{W}p")
            run = ET.SubElement(paragraph, f"{W}r")
            text_element = ET.SubElement(run, f"{W}t")
            text_element.text = paragraph_text

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _build_endnotes_part(pending_endnotes: list[tuple[str, str | tuple[str, ...]]]) -> bytes:
    root = ET.Element(f"{W}endnotes")
    separator = ET.SubElement(root, f"{W}endnote", {f"{W}type": "separator", f"{W}id": "-1"})
    separator_paragraph = ET.SubElement(separator, f"{W}p")
    separator_run = ET.SubElement(separator_paragraph, f"{W}r")
    ET.SubElement(separator_run, f"{W}separator")

    continuation = ET.SubElement(
        root,
        f"{W}endnote",
        {f"{W}type": "continuationSeparator", f"{W}id": "0"},
    )
    continuation_paragraph = ET.SubElement(continuation, f"{W}p")
    continuation_run = ET.SubElement(continuation_paragraph, f"{W}r")
    ET.SubElement(continuation_run, f"{W}continuationSeparator")

    for endnote_id, (_placeholder, text) in enumerate(pending_endnotes, start=1):
        endnote = ET.SubElement(root, f"{W}endnote", {f"{W}id": str(endnote_id)})
        for paragraph_text in _normalize_note_text(text):
            paragraph = ET.SubElement(endnote, f"{W}p")
            run = ET.SubElement(paragraph, f"{W}r")
            text_element = ET.SubElement(run, f"{W}t")
            text_element.text = paragraph_text

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _normalize_note_text(text: str | tuple[str, ...]) -> tuple[str, ...]:
    if isinstance(text, tuple):
        return text
    return (text,)
