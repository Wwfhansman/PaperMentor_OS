from pathlib import Path

from docx import Document

from papermentor_os.parsers.docx_parser import DocxPaperParser
from papermentor_os.schemas.types import Discipline, PaperStage
from tests.fixtures.review_cases import (
    ABBREVIATED_SECTION_HEADER_NOISE_CASE,
    ABSTRACT_RUNNING_HEADER_VARIATION_CASE,
    APPENDIX_VARIATION_CASE,
    APPENDIX_CONTENTS_VARIATION_CASE,
    APPENDIX_FIGURE_LIST_VARIATION_CASE,
    ANNOTATION_BLOCK_VARIATION_CASE,
    AUTHOR_INFO_VARIATION_CASE,
    BACK_MATTER_VARIATION_CASE,
    BILINGUAL_ABSTRACT_CASE,
    CAPTION_VARIATION_CASE,
    COMPLEX_TABLE_FRONT_MATTER_CASE,
    COMPLEX_CONTENTS_VARIATION_CASE,
    CONTENTS_HEADER_FOOTER_VARIATION_CASE,
    CONTENTS_FIELD_CODE_VARIATION_CASE,
    CONTENTS_VARIATION_CASE,
    COVER_PAGE_VARIATION_CASE,
    COVER_PAGE_TABLE_VARIATION_CASE,
    DECLARATION_VARIATION_CASE,
    DEPARTMENT_INFO_VARIATION_CASE,
    DOCX_ENDNOTE_OBJECT_VARIATION_CASE,
    DOCX_FOOTNOTE_OBJECT_VARIATION_CASE,
    DOCX_TABLE_FRONT_MATTER_CASE,
    EQUATION_CAPTION_VARIATION_CASE,
    ENGLISH_APPENDIX_VARIATION_CASE,
    FOOTER_FOOTNOTE_NOISE_VARIATION_CASE,
    FOOTNOTE_BODY_VARIATION_CASE,
    FRONT_MATTER_COMBO_VARIATION_CASE,
    FRONT_MATTER_MULTILINE_VARIATION_CASE,
    FRONT_MATTER_SPACING_VARIATION_CASE,
    FRONT_MATTER_TABLE_VARIATION_CASE,
    KEYWORD_VARIATION_CASE,
    METADATA_BLOCK_VARIATION_CASE,
    MIXED_NOTE_OBJECT_VARIATION_CASE,
    POST_REFERENCE_BIO_VARIATION_CASE,
    REPEATED_PARENT_SECTION_HEADER_NOISE_CASE,
    REPEATED_SECTION_HEADER_NOISE_CASE,
    REPEATED_SUBSECTION_HEADER_NOISE_CASE,
    RUNNING_ENGLISH_HEADER_FOOTER_CASE,
    RUNNING_HEADER_FOOTER_METADATA_CASE,
    TABLE_ADJACENT_NOTE_OBJECT_VARIATION_CASE,
    TEMPLATE_VARIATION_CASE,
    MULTILINE_NOTE_OBJECT_VARIATION_CASE,
    UNNUMBERED_ABBREVIATED_SECTION_HEADER_NOISE_CASE,
)
from tests.fixtures.sample_docx import build_docx_from_case


def test_docx_parser_extracts_title_abstract_sections_and_references(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("基于知识图谱的论文评审辅助系统设计")
    document.add_paragraph("摘要", style="Heading 1")
    document.add_paragraph("本文针对本科毕业论文初审反馈滞后的问题，提出一套结构化评审系统。")
    document.add_paragraph("1 绪论", style="Heading 1")
    document.add_paragraph("随着毕业论文指导工作量上升，早期诊断需求越来越突出。")
    document.add_paragraph("2 系统设计", style="Heading 1")
    document.add_paragraph("本文设计了多智能体评审流程，并通过证据账本约束结论。")
    document.add_paragraph("参考文献", style="Heading 1")
    document.add_paragraph("[1] Author. Example Paper. 2024.")
    document.save(file_path)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert paper.title == "基于知识图谱的论文评审辅助系统设计"
    assert "结构化评审系统" in paper.abstract
    assert len(paper.sections) == 2
    assert paper.sections[0].heading == "1 绪论"
    assert paper.sections[0].paragraphs[0].anchor_id == "sec-001-p-001"
    assert len(paper.references) == 1


def test_docx_parser_handles_spaced_abstract_and_reference_titles(tmp_path: Path) -> None:
    file_path = tmp_path / "template_variation.docx"
    build_docx_from_case(file_path, TEMPLATE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert "模板差异适配问题" in paper.abstract
    assert paper.sections[0].heading == "第一章 绪论"
    assert len(paper.references) == 6


def test_docx_parser_handles_cover_page_content_before_title(tmp_path: Path) -> None:
    file_path = tmp_path / "cover_page_variation.docx"
    build_docx_from_case(file_path, COVER_PAGE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert paper.title == COVER_PAGE_VARIATION_CASE.title
    assert "封面与标题布局方式" in paper.abstract
    assert paper.sections[0].heading == "1 绪论"


def test_docx_parser_handles_cover_page_table_content_before_title(tmp_path: Path) -> None:
    file_path = tmp_path / "cover_page_table_variation.docx"
    build_docx_from_case(file_path, COVER_PAGE_TABLE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert paper.title == COVER_PAGE_TABLE_VARIATION_CASE.title
    assert "封面表格模板适配问题" in paper.abstract
    assert "学号 2020123456" not in paper.abstract
    assert paper.sections[0].heading == "1 绪论"


def test_docx_parser_skips_back_matter_before_references(tmp_path: Path) -> None:
    file_path = tmp_path / "back_matter_variation.docx"
    build_docx_from_case(file_path, BACK_MATTER_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "致谢" not in headings
    assert "攻读学位期间取得的成果" not in headings
    assert "Acknowledgements" not in headings
    assert "感谢指导教师、实验室同学和参与调研的同学在论文撰写过程中的帮助。" not in section_text
    assert len(paper.sections) == 5
    assert len(paper.references) == 6


def test_docx_parser_keeps_figure_and_table_captions_inside_current_sections(tmp_path: Path) -> None:
    file_path = tmp_path / "caption_variation.docx"
    build_docx_from_case(file_path, CAPTION_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "图 1-1 论文评审系统总体架构图" not in headings
    assert "表 1-1 核心评审维度与指标说明" not in headings
    assert "Figure 2-1 Baseline parsing workflow" not in headings
    assert "Table 2-1 Error categories and examples" not in headings
    assert "图 1-1 论文评审系统总体架构图" in section_text
    assert "Figure 2-1 Baseline parsing workflow" in section_text
    assert len(paper.sections) == 5


def test_docx_parser_keeps_equation_captions_inside_current_sections(tmp_path: Path) -> None:
    file_path = tmp_path / "equation_caption_variation.docx"
    build_docx_from_case(file_path, EQUATION_CAPTION_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "式 1-1 论文评审损失函数定义" not in headings
    assert "公式 1-2 维度加权得分计算方式" not in headings
    assert "Equation 2-1 Review score aggregation" not in headings
    assert "Equation (2-2) Evidence consistency score" not in headings
    assert "式 1-1 论文评审损失函数定义" in section_text
    assert "Equation (2-2) Evidence consistency score" in section_text
    assert len(paper.sections) == 5


def test_docx_parser_keeps_annotation_blocks_inside_current_sections(tmp_path: Path) -> None:
    file_path = tmp_path / "annotation_block_variation.docx"
    build_docx_from_case(file_path, ANNOTATION_BLOCK_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "注 1-1 图片仅用于展示系统流程" not in headings
    assert "说明 1-2 指标结果采用三轮平均值" not in headings
    assert "Note 2-1 Scores are normalized" not in headings
    assert "Remark 2-2 Evidence anchors are sampled" not in headings
    assert "注 1-1 图片仅用于展示系统流程" in section_text
    assert "Remark 2-2 Evidence anchors are sampled" in section_text
    assert len(paper.sections) == 5


def test_docx_parser_filters_footer_and_footnote_noise_lines(tmp_path: Path) -> None:
    file_path = tmp_path / "footer_footnote_noise_variation.docx"
    build_docx_from_case(file_path, FOOTER_FOOTNOTE_NOISE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "- 1 -" not in headings
    assert "ii" not in headings
    assert "[1]" not in headings
    assert "①" not in headings
    assert "- 1 -" not in section_text
    assert "ii" not in section_text
    assert "[1]" not in section_text
    assert "①" not in section_text
    assert len(paper.sections) == 5


def test_docx_parser_filters_running_header_footer_metadata_rows(tmp_path: Path) -> None:
    file_path = tmp_path / "running_header_footer_metadata.docx"
    build_docx_from_case(file_path, RUNNING_HEADER_FOOTER_METADATA_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "某某大学本科毕业论文" not in headings
    assert "本科毕业论文" not in headings
    assert RUNNING_HEADER_FOOTER_METADATA_CASE.title not in section_text
    assert "学号 2020123456" not in section_text
    assert "专业 软件工程" not in section_text
    assert "指导教师 张老师" not in section_text
    assert len(paper.sections) == 5


def test_docx_parser_filters_running_english_header_footer_metadata_rows(tmp_path: Path) -> None:
    file_path = tmp_path / "running_english_header_footer.docx"
    build_docx_from_case(file_path, RUNNING_ENGLISH_HEADER_FOOTER_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "Undergraduate Thesis" not in section_text
    assert RUNNING_ENGLISH_HEADER_FOOTER_CASE.title not in section_text
    assert "Student ID 2020123456" not in section_text
    assert "Major Software Engineering" not in section_text
    assert "Bachelor Thesis" not in section_text
    assert "Advisor Prof. Zhang" not in section_text
    assert len(paper.sections) == 5


def test_docx_parser_ignores_running_abstract_headers_after_body_started(tmp_path: Path) -> None:
    file_path = tmp_path / "abstract_running_header_variation.docx"
    build_docx_from_case(file_path, ABSTRACT_RUNNING_HEADER_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "Abstract" not in {section.heading for section in paper.sections}
    assert "Abstract" not in section_text
    assert any("英文摘要页眉时仍能稳定保留后续论证内容" in text for text in section_text)
    assert "英文摘要页眉时仍能稳定保留后续论证内容" not in paper.abstract
    assert len(paper.sections) == 5


def test_docx_parser_ignores_repeated_running_section_headers(tmp_path: Path) -> None:
    file_path = tmp_path / "repeated_section_header_noise.docx"
    build_docx_from_case(file_path, REPEATED_SECTION_HEADER_NOISE_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert headings.count("1 绪论") == 1
    assert headings.count("2 已有方法比较") == 1
    assert headings.count("3 方法设计") == 1
    assert "1 绪论" not in section_text
    assert "2 已有方法比较" not in section_text
    assert "3 方法设计" not in section_text
    assert len(paper.sections) == 5


def test_docx_parser_ignores_abbreviated_section_headers_inside_deeper_subsections(tmp_path: Path) -> None:
    file_path = tmp_path / "abbreviated_section_header_noise_case.docx"
    build_docx_from_case(file_path, ABBREVIATED_SECTION_HEADER_NOISE_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert [section.heading for section in paper.sections].count("2.1 数据集构建与标注策略") == 1
    assert [section.heading for section in paper.sections].count("2.1.1 标注原则") == 1
    assert "2.1 数据集构建" not in section_text
    assert any("再次遇到已出现章节标题的缩写时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 6


def test_docx_parser_ignores_unnumbered_abbreviated_section_headers_inside_deeper_subsections(tmp_path: Path) -> None:
    file_path = tmp_path / "unnumbered_abbreviated_section_header_noise_case.docx"
    build_docx_from_case(file_path, UNNUMBERED_ABBREVIATED_SECTION_HEADER_NOISE_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert [section.heading for section in paper.sections].count("2.1 数据集构建与标注策略") == 1
    assert [section.heading for section in paper.sections].count("2.1.1 标注原则") == 1
    assert "数据集构建" not in section_text
    assert any("再次遇到已出现编号章节标题的无编号缩写时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 6


def test_docx_parser_filters_running_footnote_body_blocks(tmp_path: Path) -> None:
    file_path = tmp_path / "footnote_body_variation.docx"
    build_docx_from_case(file_path, FOOTNOTE_BODY_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "[1] 注：实验平台为课程内部服务器，仅用于教学验证。" not in section_text
    assert "① Remark: baseline scores come from the prior semester report." not in section_text
    assert any("遇到脚注正文块时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 5


def test_docx_parser_filters_docx_footnote_object_text_duplicates(tmp_path: Path) -> None:
    file_path = tmp_path / "docx_footnote_object_variation.docx"
    build_docx_from_case(file_path, DOCX_FOOTNOTE_OBJECT_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "注：实验平台为课程内部服务器，仅用于教学验证。" not in section_text
    assert any("存在真实 docx 脚注对象时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 5


def test_docx_parser_filters_docx_endnote_object_text_duplicates(tmp_path: Path) -> None:
    file_path = tmp_path / "docx_endnote_object_variation.docx"
    build_docx_from_case(file_path, DOCX_ENDNOTE_OBJECT_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "说明：附加实验参数记录在课程归档仓库中。" not in section_text
    assert any("存在真实 docx 尾注对象时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 5


def test_docx_parser_filters_mixed_note_duplicates_but_keeps_similar_body_sentence(tmp_path: Path) -> None:
    file_path = tmp_path / "mixed_note_object_variation.docx"
    build_docx_from_case(file_path, MIXED_NOTE_OBJECT_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "说明：课程内网服务器用于复现实验。" not in section_text
    assert "说明：课程镜像服务器用于补充归档。" not in section_text
    assert "说明：课程内网服务器与课程镜像服务器共同支撑实验复现。" in section_text
    assert len(paper.sections) == 5


def test_docx_parser_filters_multiline_note_duplicates_but_keeps_summary_sentence(tmp_path: Path) -> None:
    file_path = tmp_path / "multiline_note_object_variation.docx"
    build_docx_from_case(file_path, MULTILINE_NOTE_OBJECT_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "注：实验平台部署在课程内网服务器。" not in section_text
    assert "注：相关配置仅用于教学验证。" not in section_text
    assert "说明：补充日志保存在课程归档仓库。" not in section_text
    assert "说明：归档信息用于结果复核。" not in section_text
    assert "实验平台、相关配置和归档信息共同支撑教学验证与结果复核。" in section_text
    assert len(paper.sections) == 5


def test_docx_parser_filters_table_adjacent_note_duplicates_but_keeps_summary_sentence(tmp_path: Path) -> None:
    file_path = tmp_path / "table_adjacent_note_object_variation.docx"
    build_docx_from_case(file_path, TABLE_ADJACENT_NOTE_OBJECT_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "注：课程内网服务器用于复现实验。" not in section_text
    assert "说明：课程镜像服务器用于补充归档。" not in section_text
    assert "课程内网服务器与课程镜像服务器共同支撑实验复现与结果归档。" in section_text
    assert len(paper.sections) == 5


def test_docx_parser_ignores_repeated_parent_section_headers_inside_subsections(tmp_path: Path) -> None:
    file_path = tmp_path / "repeated_parent_section_header_noise.docx"
    build_docx_from_case(file_path, REPEATED_PARENT_SECTION_HEADER_NOISE_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert headings.count("2 相关工作与方法设计") == 1
    assert headings.count("2.1 数据集构建") == 1
    assert "2 相关工作与方法设计" not in section_text
    assert any("子章节内部再次遇到上级章节标题时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 5


def test_docx_parser_ignores_repeated_subsection_headers_inside_deeper_subsections(tmp_path: Path) -> None:
    file_path = tmp_path / "repeated_subsection_header_noise.docx"
    build_docx_from_case(file_path, REPEATED_SUBSECTION_HEADER_NOISE_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert headings.count("2.1 数据集构建") == 1
    assert headings.count("2.1.1 标注原则") == 1
    assert "2.1 数据集构建" not in section_text
    assert any("更细一级小节内部再次遇到子章节标题时仍能稳定保留后续论证内容" in text for text in section_text)
    assert len(paper.sections) == 6


def test_docx_parser_ignores_table_of_contents_entries_between_abstract_and_body(tmp_path: Path) -> None:
    file_path = tmp_path / "contents_variation.docx"
    build_docx_from_case(file_path, CONTENTS_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "目录" not in headings
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_preserves_bilingual_abstract_before_body(tmp_path: Path) -> None:
    file_path = tmp_path / "bilingual_abstract.docx"
    build_docx_from_case(file_path, BILINGUAL_ABSTRACT_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert "双语摘要模板适配分析问题" in paper.abstract
    assert "This paper studies bilingual abstract layouts" in paper.abstract
    assert paper.sections[0].heading == "第一章 绪论"


def test_docx_parser_skips_declaration_pages_before_body(tmp_path: Path) -> None:
    file_path = tmp_path / "declaration_variation.docx"
    build_docx_from_case(file_path, DECLARATION_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "独创性声明" not in headings
    assert "学术诚信承诺书" not in headings
    assert paper.sections[0].heading == "第一章 绪论"
    assert "本人声明所提交的毕业论文" not in paper.abstract


def test_docx_parser_ignores_complex_contents_entries_with_tabs_and_nested_numbering(tmp_path: Path) -> None:
    file_path = tmp_path / "complex_contents_variation.docx"
    build_docx_from_case(file_path, COMPLEX_CONTENTS_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "Contents" not in headings
    assert "1.1Research Background\t2" not in headings
    assert headings[0] == "1 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_ignores_contents_header_footer_noise(tmp_path: Path) -> None:
    file_path = tmp_path / "contents_header_footer_variation.docx"
    build_docx_from_case(file_path, CONTENTS_HEADER_FOOTER_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "某某大学本科毕业论文" not in headings
    assert "本科毕业论文" not in headings
    assert "- 1 -" not in headings
    assert "2" not in headings
    assert "某某大学本科毕业论文" not in section_text
    assert headings[0] == "1 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_ignores_contents_field_code_artifacts(tmp_path: Path) -> None:
    file_path = tmp_path / "contents_field_code_variation.docx"
    build_docx_from_case(file_path, CONTENTS_FIELD_CODE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "Contents" not in headings
    assert "PAGEREF _Toc482011111 \\h" not in headings
    assert 'HYPERLINK \\l "_Toc482011112"' not in headings
    assert "_Toc482011113" not in headings
    assert "TOC \\o \"1-3\" \\h \\z \\u" not in section_text
    assert headings[0] == "1 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_keeps_keyword_blocks_out_of_body_sections(tmp_path: Path) -> None:
    file_path = tmp_path / "keyword_variation.docx"
    build_docx_from_case(file_path, KEYWORD_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert "关键词：本科论文初审" in paper.abstract
    assert "Key words: thesis review" in paper.abstract
    assert "关键词：本科论文初审；多智能体评审；docx 解析；benchmark" not in {
        section.heading for section in paper.sections
    }
    assert paper.sections[0].heading == "第一章 绪论"


def test_docx_parser_skips_metadata_block_after_abstract(tmp_path: Path) -> None:
    file_path = tmp_path / "metadata_block_variation.docx"
    build_docx_from_case(file_path, METADATA_BLOCK_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert "关键词：本科论文初审" in paper.abstract
    assert "分类号：TP391.1" not in paper.abstract
    assert "学校代码：10487" not in paper.abstract
    assert "学号：2020123456" not in paper.abstract
    assert "UDC: 004.8" not in paper.abstract
    assert "分类号：TP391.1" not in {section.heading for section in paper.sections}
    assert paper.sections[0].heading == "第一章 绪论"


def test_docx_parser_skips_author_info_block_after_abstract(tmp_path: Path) -> None:
    file_path = tmp_path / "author_info_variation.docx"
    build_docx_from_case(file_path, AUTHOR_INFO_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert "关键词：本科论文初审" in paper.abstract
    assert "保密级别：公开" not in paper.abstract
    assert "作者姓名：张三" not in paper.abstract
    assert "指导教师：李老师" not in paper.abstract
    assert "作者姓名：张三" not in {section.heading for section in paper.sections}
    assert paper.sections[0].heading == "第一章 绪论"


def test_docx_parser_skips_appendix_heading_and_body_from_main_sections(tmp_path: Path) -> None:
    file_path = tmp_path / "appendix_variation.docx"
    build_docx_from_case(file_path, APPENDIX_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "附录 A" not in headings
    assert "附录 A 系统提示词样例......................21" not in headings
    assert len(paper.sections) == 5
    assert len(paper.references) == 6


def test_docx_parser_skips_english_appendix_headings_with_titles(tmp_path: Path) -> None:
    file_path = tmp_path / "english_appendix_variation.docx"
    build_docx_from_case(file_path, ENGLISH_APPENDIX_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "Appendix A Prompt Templates" not in headings
    assert "Appendix B Supplementary Results" not in headings
    assert len(paper.sections) == 5
    assert len(paper.references) == 6


def test_docx_parser_skips_department_info_block_after_abstract(tmp_path: Path) -> None:
    file_path = tmp_path / "department_info_variation.docx"
    build_docx_from_case(file_path, DEPARTMENT_INFO_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert "关键词：本科论文初审" in paper.abstract
    assert "学院：计算机科学与技术学院" not in paper.abstract
    assert "专业：软件工程" not in paper.abstract
    assert "班级：2020级1班" not in paper.abstract
    assert "作者单位：某某大学信息学院" not in paper.abstract
    assert "学院：计算机科学与技术学院" not in {section.heading for section in paper.sections}
    assert paper.sections[0].heading == "第一章 绪论"


def test_docx_parser_handles_combined_front_matter_blocks_before_body(tmp_path: Path) -> None:
    file_path = tmp_path / "front_matter_combo_variation.docx"
    build_docx_from_case(file_path, FRONT_MATTER_COMBO_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "关键词：本科论文初审" in paper.abstract
    assert "Key words: thesis review" in paper.abstract
    assert "导师组：智能系统导师组" not in paper.abstract
    assert "学位类型：工学学士" not in paper.abstract
    assert "Contents" not in headings
    assert "Appendix A Prompt Templates" not in headings
    assert "附录 B 补充图表" not in headings
    assert "图 B-1 复杂前置区模板示意图" not in section_text
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_handles_space_separated_front_matter_lines(tmp_path: Path) -> None:
    file_path = tmp_path / "front_matter_spacing_variation.docx"
    build_docx_from_case(file_path, FRONT_MATTER_SPACING_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "关键词 本科论文初审" in paper.abstract
    assert "Key words thesis review" in paper.abstract
    assert "学校代码 10487" not in paper.abstract
    assert "导师组 智能系统导师组" not in paper.abstract
    assert "学位类型 工学学士" not in paper.abstract
    assert "学校代码 10487" not in headings
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_handles_multiline_front_matter_values(tmp_path: Path) -> None:
    file_path = tmp_path / "front_matter_multiline_variation.docx"
    build_docx_from_case(file_path, FRONT_MATTER_MULTILINE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "本科论文初审；跨行前置区" in paper.abstract
    assert "thesis review; multiline metadata; parser stability" in paper.abstract
    assert "10487" not in headings
    assert "智能系统导师组" not in headings
    assert "工学学士" not in headings
    assert "10487" not in section_text
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_skips_appendix_contents_page_and_following_appendices(tmp_path: Path) -> None:
    file_path = tmp_path / "appendix_contents_variation.docx"
    build_docx_from_case(file_path, APPENDIX_CONTENTS_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "附录目录" not in headings
    assert "附录 A 访谈提纲" not in headings
    assert "附录 B 问卷样例" not in headings
    assert "附录图目录....................................27" not in section_text
    assert len(paper.sections) == 5
    assert len(paper.references) == 6


def test_docx_parser_handles_tabular_front_matter_lines(tmp_path: Path) -> None:
    file_path = tmp_path / "front_matter_table_variation.docx"
    build_docx_from_case(file_path, FRONT_MATTER_TABLE_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "关键词 本科论文初审" in paper.abstract
    assert "Key words thesis review" in paper.abstract
    assert "学校代码 10487" not in paper.abstract
    assert "导师组 智能系统导师组" not in paper.abstract
    assert "学校代码 10487" not in headings
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_skips_appendix_figure_and_table_list_pages(tmp_path: Path) -> None:
    file_path = tmp_path / "appendix_figure_list_variation.docx"
    build_docx_from_case(file_path, APPENDIX_FIGURE_LIST_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "附录图目录" not in headings
    assert "List of Appendix Tables" not in headings
    assert "附录 A 访谈材料" not in headings
    assert "Appendix B Supplementary Questionnaire" not in headings
    assert "图 A-1 访谈流程图................................21" not in section_text
    assert len(paper.sections) == 5
    assert len(paper.references) == 6


def test_docx_parser_extracts_front_matter_from_docx_tables(tmp_path: Path) -> None:
    file_path = tmp_path / "docx_table_front_matter.docx"
    build_docx_from_case(file_path, DOCX_TABLE_FRONT_MATTER_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]

    assert "关键词 本科论文初审" in paper.abstract
    assert "Key words thesis review" in paper.abstract
    assert "学校代码 10487" not in paper.abstract
    assert "导师组 智能系统导师组" not in paper.abstract
    assert "学校代码 10487" not in headings
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_extracts_front_matter_from_merged_and_nested_docx_tables(tmp_path: Path) -> None:
    file_path = tmp_path / "complex_table_front_matter.docx"
    build_docx_from_case(file_path, COMPLEX_TABLE_FRONT_MATTER_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    headings = [section.heading for section in paper.sections]
    section_text = {
        paragraph.text
        for section in paper.sections
        for paragraph in section.paragraphs
    }

    assert "关键词 本科论文初审" in paper.abstract
    assert "Key words thesis review" in paper.abstract
    assert "学校代码 10487" not in paper.abstract
    assert "学校代码 学校代码 10487" not in paper.abstract
    assert "导师组 智能系统导师组 学位类型 工学学士 学位授予单位 某某大学" not in paper.abstract
    assert "学校代码 10487" not in headings
    assert "导师组 智能系统导师组 学位类型 工学学士 学位授予单位 某某大学" not in section_text
    assert headings[0] == "第一章 绪论"
    assert len(paper.sections) == 5


def test_docx_parser_stops_references_before_author_biography_pages(tmp_path: Path) -> None:
    file_path = tmp_path / "post_reference_bio_variation.docx"
    build_docx_from_case(file_path, POST_REFERENCE_BIO_VARIATION_CASE)

    parser = DocxPaperParser()
    paper = parser.parse_file(
        file_path,
        stage=PaperStage.DRAFT,
        discipline=Discipline.COMPUTER_SCIENCE,
    )

    assert len(paper.references) == 6
    assert all("作者简介" not in reference.raw_text for reference in paper.references)
    assert all("Author Biography" not in reference.raw_text for reference in paper.references)
    assert len(paper.sections) == 5
